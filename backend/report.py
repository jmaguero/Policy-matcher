import os
import re
from datetime import datetime
from pathlib import Path

import openpyxl
from docx import Document

OUTPUTS_DIR = Path(__file__).parent / "outputs"

# Docker sets TEMPLATE_PATH=/app/documents/template_report.docx via env var.
# Local dev falls back to the documents/ folder next to the backend/ directory.
TEMPLATE_PATH = Path(
    os.getenv("TEMPLATE_PATH")
    or str(Path(__file__).parent.parent / "documents" / "template_report.docx")
)

# LLM05: output sanitization limits (data originates from LLM via the xlsx)
MAX_TEXT_CHARS = 5_000
MAX_BULLET_ITEMS = 10
MAX_BULLET_CHARS = 2_000

# Control characters to strip from LLM-generated text before inserting into docx.
# Null bytes and non-printable chars (except tab and newline) corrupt OOXML.
_CTRL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _sanitize_text(value: object, max_len: int = MAX_TEXT_CHARS) -> str:
    """LLM05: coerce to string, strip control chars, cap length."""
    if value is None:
        return ""
    s = str(value)
    s = _CTRL_RE.sub("", s)
    return s[:max_len]


def _resolve_input(filename: str) -> Path:
    """Resolve filename to OUTPUTS_DIR; reject path traversal."""
    resolved = (OUTPUTS_DIR / filename).resolve()
    if not resolved.is_relative_to(OUTPUTS_DIR.resolve()):
        raise ValueError("Invalid input filename")
    if not resolved.exists():
        raise FileNotFoundError(f"File not found in outputs: {filename}")
    return resolved


def _read_xlsx_rows(xlsx_path: Path) -> list[dict]:
    try:
        wb = openpyxl.load_workbook(xlsx_path)
    except Exception:
        raise ValueError("Could not open the input XLSX file")
    ws = wb.active
    headers = [cell.value for cell in ws[1]]
    return [
        dict(zip(headers, row)) for row in ws.iter_rows(min_row=2, values_only=True)
    ]


def _parse_bullet_items(raw: object) -> list[str]:
    """Split the newline-joined rewritten_suggestions back into individual items."""
    if raw is None:
        return []
    items = str(raw).split("\n")
    # Cap count and length (LLM05)
    items = [
        _sanitize_text(item, MAX_BULLET_CHARS) for item in items if str(item).strip()
    ]
    return items[:MAX_BULLET_ITEMS]


def _add_bullet(doc: Document, text: str) -> None:
    """Add a bullet paragraph using List Paragraph style + unicode bullet."""
    p = doc.add_paragraph(style="List Paragraph")
    p.add_run(f"\u2022 {text}")


def generate_report(input_xlsx_filename: str) -> dict:
    OUTPUTS_DIR.mkdir(exist_ok=True)

    # Validate template exists (server-side, hardcoded — not user-supplied)
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Report template not found: {TEMPLATE_PATH}")

    input_path = _resolve_input(input_xlsx_filename)
    rows = _read_xlsx_rows(input_path)

    try:
        doc = Document(str(TEMPLATE_PATH))
    except Exception:
        raise ValueError("Could not open the report template")

    # Track whether we've written at least one section
    wrote_any = False

    processed_rows = []
    for row in rows:
        suggestions_raw = (row.get("suggestions") or "").strip()
        if not suggestions_raw:
            continue  # only report rows that have suggestions

        # LLM05: sanitize all LLM-origin fields before inserting into docx
        category = _sanitize_text(row.get("category"))
        title = _sanitize_text(row.get("title"))
        control = _sanitize_text(row.get("control"))
        bullet_items = _parse_bullet_items(row.get("rewritten_suggestions"))

        # Store processed row for sorting/grouping
        processed_rows.append(
            {
                "category": category,
                "title": title,
                "control": control,
                "items": bullet_items,
                "id": row.get("id", ""),
            }
        )

    # Sort rows by category to ensure grouping works
    processed_rows.sort(key=lambda x: (x["category"] or "", x["id"] or ""))

    last_category = None

    for row in processed_rows:
        if not wrote_any:
            # Page break before first finding — keeps cover page clean
            doc.add_page_break()
            wrote_any = True

        if row["category"] != last_category:
            doc.add_heading(row["category"], level=2)
            last_category = row["category"]

        doc.add_heading(row["title"], level=3)
        doc.add_paragraph(row["control"])

        for item in row["items"]:
            _add_bullet(doc, item)

    # Derive the client+standard part from input filename to carry it forward.
    # Input stem format: <YYYYMMDD>_<HHMMSS>_<client>_<standard>
    stem = Path(input_xlsx_filename).stem
    parts = stem.split("_", 2)
    remainder = parts[2] if len(parts) > 2 else stem

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_out_path = OUTPUTS_DIR / f"{ts}_{remainder}.docx"
    doc.save(str(docx_out_path))

    return {
        "docx_file": docx_out_path.name,
        "rows_reported": wrote_any,
    }
